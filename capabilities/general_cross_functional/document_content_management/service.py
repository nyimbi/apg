"""
Document Content Management - Comprehensive Service Layer

Enterprise-grade document management service providing business logic for document operations,
workflow management, version control, permissions, content analysis, and collaboration.

Copyright © 2025 Datacraft
Author: Nyimbi Odero <nyimbi@gmail.com>
Website: www.datacraft.co.ke
"""

import asyncio
import hashlib
import mimetypes
import os
import shutil
from datetime import datetime, date, timedelta
from decimal import Decimal
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union, BinaryIO
from uuid import UUID
import json
import tempfile
import zipfile

from pydantic import ValidationError
import aiofiles
import aiofiles.os
from PIL import Image, ImageDraw, ImageFont
import magic
from langdetect import detect
import nltk
from nltk.sentiment import SentimentIntensityAnalyzer
import textstat
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import pytesseract
import cv2
import numpy as np
from moviepy.editor import VideoFileClip
import fitz  # PyMuPDF
import docx
import openpyxl
from lxml import etree
import markdown
import bleach

from .models import *
from ..integration_api_management.service import IntegrationAPIService
from ..advanced_analytics_platform.service import AdvancedAnalyticsService


class DocumentContentManagementError(Exception):
	"""Base exception for DCM operations"""
	pass


class DocumentNotFoundError(DocumentContentManagementError):
	"""Document not found error"""
	pass


class PermissionDeniedError(DocumentContentManagementError):
	"""Permission denied error"""
	pass


class WorkflowError(DocumentContentManagementError):
	"""Workflow operation error"""
	pass


class DocumentContentManagementService:
	"""
	Comprehensive Document Content Management service providing enterprise-grade
	document management, content collaboration, and knowledge management capabilities.
	"""
	
	def __init__(self, 
				 storage_root: str = "/data/dcm",
				 max_file_size: int = 500 * 1024 * 1024,  # 500MB
				 supported_formats: Optional[List[str]] = None,
				 enable_ai_analysis: bool = True):
		self.storage_root = Path(storage_root)
		self.max_file_size = max_file_size
		self.supported_formats = supported_formats or self._default_supported_formats()
		self.enable_ai_analysis = enable_ai_analysis
		
		# Initialize storage structure
		self._init_storage_structure()
		
		# Initialize NLP tools
		if enable_ai_analysis:
			self._init_nlp_tools()
		
		# Initialize external services
		try:
			self.integration_service = IntegrationAPIService()
			self.analytics_service = AdvancedAnalyticsService()
		except:
			# Services may not be available in all environments
			self.integration_service = None
			self.analytics_service = None
		
		# Cache for frequently accessed data
		self._permission_cache = {}
		self._workflow_cache = {}
		
	def _default_supported_formats(self) -> List[str]:
		"""Default supported file formats"""
		return [
			'pdf', 'docx', 'doc', 'xlsx', 'xls', 'pptx', 'ppt',
			'txt', 'rtf', 'html', 'htm', 'md', 'csv', 'json', 'xml',
			'jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'svg',
			'mp4', 'avi', 'mov', 'wmv', 'flv', 'webm',
			'mp3', 'wav', 'flac', 'aac', 'ogg',
			'zip', 'rar', '7z', 'tar', 'gz',
			'dwg', 'dxf', 'step', 'iges'
		]
	
	def _init_storage_structure(self):
		"""Initialize storage directory structure"""
		directories = [
			'documents', 'versions', 'previews', 'thumbnails',
			'assets', 'temp', 'exports', 'archives', 'quarantine'
		]
		
		for directory in directories:
			(self.storage_root / directory).mkdir(parents=True, exist_ok=True)
	
	def _init_nlp_tools(self):
		"""Initialize NLP and AI analysis tools"""
		try:
			nltk.download('vader_lexicon', quiet=True)
			nltk.download('punkt', quiet=True)
			nltk.download('stopwords', quiet=True)
			self.sentiment_analyzer = SentimentIntensityAnalyzer()
			self.vectorizer = TfidfVectorizer(max_features=1000, stop_words='english')
		except Exception as e:
			print(f"Warning: Failed to initialize NLP tools: {e}")
			self.enable_ai_analysis = False
	
	# Core Document Management
	async def create_document(self, 
							  document_data: Dict[str, Any],
							  file_content: Optional[BinaryIO] = None,
							  tenant_id: str,
							  user_id: str) -> DCMDocument:
		"""Create a new document with optional file content"""
		
		# Validate input data
		if not document_data.get('name'):
			raise DocumentContentManagementError("Document name is required")
		
		# Generate document ID and storage path
		document_id = uuid7str()
		
		# Process file if provided
		file_info = {}
		if file_content:
			file_info = await self._process_uploaded_file(file_content, document_id)
		
		# Create document model
		document = DCMDocument(
			id=document_id,
			tenant_id=tenant_id,
			created_by=user_id,
			updated_by=user_id,
			name=document_data['name'],
			title=document_data.get('title', document_data['name']),
			description=document_data.get('description'),
			folder_id=document_data.get('folder_id'),
			document_type=document_data.get('document_type', DCMDocumentType.TEXT_DOCUMENT),
			**file_info
		)
		
		# Perform content analysis if enabled
		if self.enable_ai_analysis and file_info:
			await self._analyze_document_content(document)
		
		# Create search index
		await self._create_search_index(document)
		
		# Log document creation
		await self._log_audit_event(
			resource_id=document.id,
			resource_type="document",
			action="create",
			action_category="modify",
			user_id=user_id,
			tenant_id=tenant_id,
			success=True
		)
		
		# Send notification
		await self._send_notification(
			notification_type=DCMNotificationType.DOCUMENT_CREATED,
			recipient_id=user_id,
			title=f"Document '{document.name}' created",
			message=f"Document '{document.name}' has been successfully created",
			document_id=document.id,
			tenant_id=tenant_id
		)
		
		return document
	
	async def get_document(self, 
						   document_id: str,
						   tenant_id: str,
						   user_id: str,
						   include_content: bool = False) -> DCMDocument:
		"""Retrieve a document with permission checking"""
		
		# Check if document exists (mock database query)
		document = await self._get_document_from_db(document_id, tenant_id)
		if not document:
			raise DocumentNotFoundError(f"Document {document_id} not found")
		
		# Check permissions
		if not await self._check_permission(document_id, user_id, DCMPermissionLevel.READ):
			raise PermissionDeniedError("Insufficient permissions to view document")
		
		# Update view statistics
		await self._update_document_stats(document_id, 'view_count')
		document.last_viewed_at = datetime.utcnow()
		
		# Load content if requested
		if include_content and document.storage_path:
			document.extracted_text = await self._load_document_content(document.storage_path)
		
		# Log access
		await self._log_audit_event(
			resource_id=document_id,
			resource_type="document",
			action="view",
			action_category="access",
			user_id=user_id,
			tenant_id=tenant_id,
			success=True
		)
		
		return document
	
	async def update_document(self,
							  document_id: str,
							  update_data: Dict[str, Any],
							  tenant_id: str,
							  user_id: str,
							  file_content: Optional[BinaryIO] = None) -> DCMDocument:
		"""Update document metadata and optionally content"""
		
		# Get existing document
		document = await self.get_document(document_id, tenant_id, user_id)
		
		# Check write permissions
		if not await self._check_permission(document_id, user_id, DCMPermissionLevel.WRITE):
			raise PermissionDeniedError("Insufficient permissions to modify document")
		
		# Check if document is locked
		if document.is_locked and document.locked_by != user_id:
			raise DocumentContentManagementError(f"Document is locked by another user")
		
		# Store old values for audit
		old_values = document.dict()
		
		# Handle file content update
		if file_content:
			# Create new version
			await self._create_document_version(document, user_id)
			
			# Process new file
			file_info = await self._process_uploaded_file(file_content, document_id)
			update_data.update(file_info)
			
			# Increment version
			document.major_version += 1
			document.minor_version = 0
			document.version_number = f"{document.major_version}.{document.minor_version}"
		
		# Update document fields
		for field, value in update_data.items():
			if hasattr(document, field):
				setattr(document, field, value)
		
		document.updated_by = user_id
		document.updated_at = datetime.utcnow()
		
		# Re-analyze content if file was updated
		if file_content and self.enable_ai_analysis:
			await self._analyze_document_content(document)
		
		# Update search index
		await self._update_search_index(document)
		
		# Log update
		await self._log_audit_event(
			resource_id=document_id,
			resource_type="document",
			action="update",
			action_category="modify",
			user_id=user_id,
			tenant_id=tenant_id,
			old_values=old_values,
			new_values=document.dict(),
			success=True
		)
		
		# Send notification
		await self._send_notification(
			notification_type=DCMNotificationType.DOCUMENT_UPDATED,
			recipient_id=user_id,
			title=f"Document '{document.name}' updated",
			message=f"Document '{document.name}' has been successfully updated",
			document_id=document.id,
			tenant_id=tenant_id
		)
		
		return document
	
	async def delete_document(self,
							  document_id: str,
							  tenant_id: str,
							  user_id: str,
							  permanent: bool = False) -> bool:
		"""Delete or archive a document"""
		
		# Get document
		document = await self.get_document(document_id, tenant_id, user_id)
		
		# Check delete permissions
		if not await self._check_permission(document_id, user_id, DCMPermissionLevel.DELETE):
			raise PermissionDeniedError("Insufficient permissions to delete document")
		
		# Check retention policy
		if not permanent:
			retention_policy = await self._get_retention_policy(document)
			if retention_policy and not await self._can_delete_per_policy(document, retention_policy):
				raise DocumentContentManagementError("Document cannot be deleted due to retention policy")
		
		if permanent:
			# Permanent deletion
			await self._permanently_delete_document(document)
		else:
			# Soft delete (archive)
			document.status = DCMDocumentStatus.ARCHIVED
			document.updated_by = user_id
			document.updated_at = datetime.utcnow()
		
		# Remove from search index
		await self._remove_from_search_index(document_id)
		
		# Log deletion
		await self._log_audit_event(
			resource_id=document_id,
			resource_type="document",
			action="delete" if permanent else "archive",
			action_category="modify",
			user_id=user_id,
			tenant_id=tenant_id,
			success=True
		)
		
		return True
	
	# Version Management
	async def create_document_version(self,
									  document_id: str,
									  change_description: str,
									  tenant_id: str,
									  user_id: str,
									  file_content: Optional[BinaryIO] = None) -> DCMDocumentVersion:
		"""Create a new version of a document"""
		
		# Get current document
		document = await self.get_document(document_id, tenant_id, user_id)
		
		# Check write permissions
		if not await self._check_permission(document_id, user_id, DCMPermissionLevel.WRITE):
			raise PermissionDeniedError("Insufficient permissions to create version")
		
		# Create version record
		version = await self._create_document_version(document, user_id, change_description)
		
		# Handle new file content
		if file_content:
			file_info = await self._process_uploaded_file(file_content, f"{document_id}_v{version.version_number}")
			version.file_name = file_info['file_name']
			version.file_size = file_info['file_size']
			version.file_hash = file_info['file_hash']
			version.storage_path = file_info['storage_path']
			
			# Update main document
			document.version_number = version.version_number
			document.major_version = version.major_version
			document.minor_version = version.minor_version
			document.file_size = file_info['file_size']
			document.file_hash = file_info['file_hash']
			document.storage_path = file_info['storage_path']
			document.updated_by = user_id
			document.updated_at = datetime.utcnow()
		
		return version
	
	async def get_document_versions(self,
									document_id: str,
									tenant_id: str,
									user_id: str) -> List[DCMDocumentVersion]:
		"""Get all versions of a document"""
		
		# Check read permissions
		if not await self._check_permission(document_id, user_id, DCMPermissionLevel.READ):
			raise PermissionDeniedError("Insufficient permissions to view document versions")
		
		# Mock database query for versions
		versions = await self._get_document_versions_from_db(document_id, tenant_id)
		
		return sorted(versions, key=lambda v: v.created_at, reverse=True)
	
	async def restore_document_version(self,
									   document_id: str,
									   version_id: str,
									   tenant_id: str,
									   user_id: str) -> DCMDocument:
		"""Restore a document to a specific version"""
		
		# Check write permissions
		if not await self._check_permission(document_id, user_id, DCMPermissionLevel.WRITE):
			raise PermissionDeniedError("Insufficient permissions to restore document version")
		
		# Get version and document
		version = await self._get_document_version_from_db(version_id, tenant_id)
		document = await self.get_document(document_id, tenant_id, user_id)
		
		if not version:
			raise DocumentNotFoundError(f"Version {version_id} not found")
		
		# Create new version from current state before restore
		await self._create_document_version(document, user_id, f"Before restore to version {version.version_number}")
		
		# Restore document to version state
		document.file_name = version.file_name
		document.file_size = version.file_size
		document.file_hash = version.file_hash
		document.storage_path = version.storage_path
		document.updated_by = user_id
		document.updated_at = datetime.utcnow()
		
		# Update version statistics
		version.restoration_count += 1
		
		return document
	
	# Permission Management
	async def set_document_permission(self,
									  document_id: str,
									  subject_id: str,
									  subject_type: str,
									  permission_level: DCMPermissionLevel,
									  tenant_id: str,
									  user_id: str,
									  **kwargs) -> DCMDocumentPermission:
		"""Set permissions for a document"""
		
		# Check admin permissions
		if not await self._check_permission(document_id, user_id, DCMPermissionLevel.ADMIN):
			raise PermissionDeniedError("Insufficient permissions to manage document permissions")
		
		permission = DCMDocumentPermission(
			tenant_id=tenant_id,
			created_by=user_id,
			updated_by=user_id,
			resource_id=document_id,
			resource_type="document",
			subject_id=subject_id,
			subject_type=subject_type,
			permission_level=permission_level,
			granted_by=user_id,
			**self._calculate_permission_flags(permission_level),
			**kwargs
		)
		
		# Clear permission cache
		self._permission_cache.pop(f"{document_id}_{subject_id}", None)
		
		return permission
	
	async def check_document_permission(self,
										document_id: str,
										user_id: str,
										required_permission: DCMPermissionLevel,
										tenant_id: str) -> bool:
		"""Check if user has required permission for document"""
		
		return await self._check_permission(document_id, user_id, required_permission)
	
	async def get_document_permissions(self,
									   document_id: str,
									   tenant_id: str,
									   user_id: str) -> List[DCMDocumentPermission]:
		"""Get all permissions for a document"""
		
		# Check admin permissions
		if not await self._check_permission(document_id, user_id, DCMPermissionLevel.ADMIN):
			raise PermissionDeniedError("Insufficient permissions to view document permissions")
		
		return await self._get_document_permissions_from_db(document_id, tenant_id)
	
	# Workflow Management
	async def start_workflow(self,
							 document_id: str,
							 workflow_id: str,
							 tenant_id: str,
							 user_id: str,
							 variables: Optional[Dict[str, Any]] = None) -> DCMWorkflowInstance:
		"""Start a workflow for a document"""
		
		# Get workflow template
		workflow = await self._get_workflow_from_db(workflow_id, tenant_id)
		if not workflow:
			raise WorkflowError(f"Workflow {workflow_id} not found")
		
		# Check permissions
		if not await self._check_permission(document_id, user_id, DCMPermissionLevel.WRITE):
			raise PermissionDeniedError("Insufficient permissions to start workflow")
		
		# Create workflow instance
		instance = DCMWorkflowInstance(
			tenant_id=tenant_id,
			created_by=user_id,
			updated_by=user_id,
			workflow_id=workflow_id,
			document_id=document_id,
			instance_name=f"{workflow.name} - {datetime.now().strftime('%Y%m%d_%H%M%S')}",
			started_by=user_id,
			workflow_variables=variables or {},
		)
		
		# Initialize first step
		await self._initialize_workflow_steps(instance, workflow)
		
		# Update workflow usage statistics
		workflow.usage_count += 1
		
		return instance
	
	async def complete_workflow_step(self,
									 instance_id: str,
									 step_id: str,
									 decision: str,
									 comments: Optional[str],
									 tenant_id: str,
									 user_id: str) -> DCMWorkflowStep:
		"""Complete a workflow step"""
		
		# Get workflow step
		step = await self._get_workflow_step_from_db(step_id, tenant_id)
		if not step:
			raise WorkflowError(f"Workflow step {step_id} not found")
		
		# Check if user is assigned to step
		if user_id not in step.assigned_to:
			raise PermissionDeniedError("User not assigned to this workflow step")
		
		# Update step
		step.status = DCMWorkflowStatus.COMPLETED
		step.decision = decision
		step.decision_by = user_id
		step.decision_at = datetime.utcnow()
		step.completed_at = datetime.utcnow()
		step.duration_hours = (step.completed_at - step.started_at).total_seconds() / 3600
		
		if comments:
			# Create comment
			comment = await self._create_workflow_comment(step, user_id, comments, tenant_id)
			step.comments.append(comment.id)
		
		# Process next step
		instance = await self._get_workflow_instance_from_db(instance_id, tenant_id)
		await self._process_next_workflow_step(instance, step)
		
		return step
	
	async def get_workflow_status(self,
								  instance_id: str,
								  tenant_id: str,
								  user_id: str) -> Dict[str, Any]:
		"""Get comprehensive workflow status"""
		
		instance = await self._get_workflow_instance_from_db(instance_id, tenant_id)
		if not instance:
			raise WorkflowError(f"Workflow instance {instance_id} not found")
		
		# Check permissions
		if not await self._check_permission(instance.document_id, user_id, DCMPermissionLevel.READ):
			raise PermissionDeniedError("Insufficient permissions to view workflow status")
		
		# Get all steps
		steps = await self._get_workflow_steps_from_db(instance_id, tenant_id)
		
		return {
			'instance': instance,
			'steps': steps,
			'progress_percentage': self._calculate_workflow_progress(instance, steps),
			'estimated_completion': self._estimate_workflow_completion(instance, steps),
			'current_bottlenecks': self._identify_workflow_bottlenecks(steps)
		}
	
	# Comment and Collaboration
	async def add_comment(self,
						  document_id: str,
						  comment_text: str,
						  tenant_id: str,
						  user_id: str,
						  parent_comment_id: Optional[str] = None,
						  **kwargs) -> DCMComment:
		"""Add a comment to a document"""
		
		# Check comment permissions
		if not await self._check_permission(document_id, user_id, DCMPermissionLevel.READ):
			raise PermissionDeniedError("Insufficient permissions to comment on document")
		
		comment = DCMComment(
			tenant_id=tenant_id,
			created_by=user_id,
			updated_by=user_id,
			document_id=document_id,
			parent_comment_id=parent_comment_id,
			comment_text=comment_text,
			**kwargs
		)
		
		# Process mentions
		mentions = self._extract_mentions(comment_text)
		comment.mention_users = mentions
		
		# Send notifications for mentions
		for mentioned_user in mentions:
			await self._send_notification(
				notification_type=DCMNotificationType.COMMENT_ADDED,
				recipient_id=mentioned_user,
				title=f"You were mentioned in a comment",
				message=f"You were mentioned in a comment on document '{await self._get_document_name(document_id)}'",
				document_id=document_id,
				comment_id=comment.id,
				tenant_id=tenant_id
			)
		
		return comment
	
	async def get_document_comments(self,
									document_id: str,
									tenant_id: str,
									user_id: str,
									include_resolved: bool = True) -> List[DCMComment]:
		"""Get all comments for a document"""
		
		# Check read permissions
		if not await self._check_permission(document_id, user_id, DCMPermissionLevel.READ):
			raise PermissionDeniedError("Insufficient permissions to view document comments")
		
		comments = await self._get_document_comments_from_db(document_id, tenant_id)
		
		if not include_resolved:
			comments = [c for c in comments if not c.is_resolved]
		
		return comments
	
	async def resolve_comment(self,
							  comment_id: str,
							  resolution_note: Optional[str],
							  tenant_id: str,
							  user_id: str) -> DCMComment:
		"""Resolve a comment"""
		
		comment = await self._get_comment_from_db(comment_id, tenant_id)
		if not comment:
			raise DocumentNotFoundError(f"Comment {comment_id} not found")
		
		# Check permissions (comment author or document admin)
		if comment.created_by != user_id and not await self._check_permission(comment.document_id, user_id, DCMPermissionLevel.ADMIN):
			raise PermissionDeniedError("Insufficient permissions to resolve comment")
		
		comment.is_resolved = True
		comment.resolved_by = user_id
		comment.resolved_at = datetime.utcnow()
		comment.resolution_note = resolution_note
		
		return comment
	
	# Search and Discovery
	async def search_documents(self,
							   query: str,
							   tenant_id: str,
							   user_id: str,
							   filters: Optional[Dict[str, Any]] = None,
							   limit: int = 50,
							   offset: int = 0) -> Dict[str, Any]:
		"""Search documents with full-text search and filtering"""
		
		# Parse search query
		search_terms = self._parse_search_query(query)
		
		# Build search filters
		search_filters = self._build_search_filters(filters or {}, tenant_id)
		
		# Perform search
		results = await self._execute_document_search(
			terms=search_terms,
			filters=search_filters,
			user_id=user_id,
			limit=limit,
			offset=offset
		)
		
		# Filter results by permissions
		authorized_results = []
		for result in results['documents']:
			if await self._check_permission(result['id'], user_id, DCMPermissionLevel.READ):
				authorized_results.append(result)
		
		# Calculate relevance scores
		scored_results = self._calculate_search_relevance(authorized_results, search_terms)
		
		return {
			'documents': scored_results,
			'total_count': len(authorized_results),
			'search_time_ms': results.get('search_time_ms', 0),
			'suggestions': await self._generate_search_suggestions(query, tenant_id)
		}
	
	async def get_document_recommendations(self,
										   document_id: str,
										   tenant_id: str,
										   user_id: str,
										   count: int = 10) -> List[Dict[str, Any]]:
		"""Get document recommendations based on content similarity"""
		
		# Get document content
		document = await self.get_document(document_id, tenant_id, user_id, include_content=True)
		
		if not document.extracted_text:
			return []
		
		# Find similar documents
		similar_docs = await self._find_similar_documents(
			content=document.extracted_text,
			document_id=document_id,
			tenant_id=tenant_id,
			user_id=user_id,
			count=count
		)
		
		return similar_docs
	
	# Content Analysis and AI
	async def analyze_document_content(self,
									   document_id: str,
									   tenant_id: str,
									   user_id: str,
									   analysis_types: Optional[List[str]] = None) -> Dict[str, Any]:
		"""Perform AI-powered content analysis"""
		
		if not self.enable_ai_analysis:
			raise DocumentContentManagementError("AI analysis is not enabled")
		
		# Get document
		document = await self.get_document(document_id, tenant_id, user_id, include_content=True)
		
		analysis_types = analysis_types or ['sentiment', 'keywords', 'summary', 'classification']
		results = {}
		
		if 'sentiment' in analysis_types:
			results['sentiment'] = await self._analyze_sentiment(document.extracted_text)
		
		if 'keywords' in analysis_types:
			results['keywords'] = await self._extract_keywords(document.extracted_text)
		
		if 'summary' in analysis_types:
			results['summary'] = await self._generate_summary(document.extracted_text)
		
		if 'classification' in analysis_types:
			results['classification'] = await self._classify_content(document.extracted_text)
		
		if 'entities' in analysis_types:
			results['entities'] = await self._extract_entities(document.extracted_text)
		
		if 'topics' in analysis_types:
			results['topics'] = await self._extract_topics(document.extracted_text)
		
		# Update document with analysis results
		document.ai_tags = results.get('keywords', [])
		document.sentiment_score = results.get('sentiment', {}).get('compound', 0.0)
		document.content_summary = results.get('summary', '')
		
		return results
	
	# Private Helper Methods
	async def _process_uploaded_file(self, file_content: BinaryIO, document_id: str) -> Dict[str, Any]:
		"""Process uploaded file and extract metadata"""
		
		# Create temporary file
		temp_path = self.storage_root / 'temp' / f"{document_id}_{uuid7str()}"
		
		async with aiofiles.open(temp_path, 'wb') as temp_file:
			content = file_content.read()
			await temp_file.write(content)
		
		try:
			# Get file information
			file_size = len(content)
			if file_size > self.max_file_size:
				raise DocumentContentManagementError(f"File size ({file_size}) exceeds maximum allowed ({self.max_file_size})")
			
			# Calculate hash
			file_hash = hashlib.sha256(content).hexdigest()
			
			# Detect MIME type
			mime_type = magic.from_buffer(content, mime=True)
			
			# Get file extension
			file_extension = Path(file_content.name if hasattr(file_content, 'name') else 'unknown').suffix.lower()
			
			# Validate format
			if file_extension.lstrip('.') not in self.supported_formats:
				raise DocumentContentManagementError(f"Unsupported file format: {file_extension}")
			
			# Move to permanent storage
			storage_path = self.storage_root / 'documents' / f"{document_id}{file_extension}"
			shutil.move(str(temp_path), str(storage_path))
			
			# Extract text content
			extracted_text = await self._extract_text_content(storage_path, mime_type)
			
			return {
				'file_name': getattr(file_content, 'name', f"document{file_extension}"),
				'file_size': file_size,
				'file_hash': file_hash,
				'mime_type': mime_type,
				'storage_path': str(storage_path),
				'content_format': DCMContentFormat(file_extension.lstrip('.')),
				'extracted_text': extracted_text
			}
		
		finally:
			# Clean up temporary file
			if temp_path.exists():
				temp_path.unlink()
	
	async def _extract_text_content(self, file_path: Path, mime_type: str) -> Optional[str]:
		"""Extract text content from various file formats"""
		
		try:
			if mime_type.startswith('text/'):
				async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
					return await f.read()
			
			elif mime_type == 'application/pdf':
				return self._extract_pdf_text(file_path)
			
			elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
				return self._extract_docx_text(file_path)
			
			elif mime_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
				return self._extract_xlsx_text(file_path)
			
			elif mime_type.startswith('image/'):
				return await self._extract_image_text(file_path)
			
			else:
				return None
		
		except Exception as e:
			print(f"Failed to extract text from {file_path}: {e}")
			return None
	
	def _extract_pdf_text(self, file_path: Path) -> str:
		"""Extract text from PDF files"""
		try:
			doc = fitz.open(str(file_path))
			text = ""
			for page in doc:
				text += page.get_text()
			doc.close()
			return text
		except Exception as e:
			print(f"Failed to extract PDF text: {e}")
			return ""
	
	def _extract_docx_text(self, file_path: Path) -> str:
		"""Extract text from DOCX files"""
		try:
			doc = docx.Document(str(file_path))
			text = ""
			for paragraph in doc.paragraphs:
				text += paragraph.text + "\n"
			return text
		except Exception as e:
			print(f"Failed to extract DOCX text: {e}")
			return ""
	
	def _extract_xlsx_text(self, file_path: Path) -> str:
		"""Extract text from XLSX files"""
		try:
			workbook = openpyxl.load_workbook(str(file_path))
			text = ""
			for sheet_name in workbook.sheetnames:
				sheet = workbook[sheet_name]
				for row in sheet.iter_rows():
					for cell in row:
						if cell.value:
							text += str(cell.value) + " "
				text += "\n"
			return text
		except Exception as e:
			print(f"Failed to extract XLSX text: {e}")
			return ""
	
	async def _extract_image_text(self, file_path: Path) -> str:
		"""Extract text from images using OCR"""
		try:
			image = cv2.imread(str(file_path))
			text = pytesseract.image_to_string(image)
			return text
		except Exception as e:
			print(f"Failed to extract image text: {e}")
			return ""
	
	async def _analyze_document_content(self, document: DCMDocument):
		"""Perform AI analysis on document content"""
		
		if not document.extracted_text or not self.enable_ai_analysis:
			return
		
		try:
			# Sentiment analysis
			sentiment = self.sentiment_analyzer.polarity_scores(document.extracted_text)
			document.sentiment_score = sentiment['compound']
			
			# Language detection
			try:
				document.language = detect(document.extracted_text)
			except:
				document.language = 'en'
			
			# Extract keywords using TF-IDF
			keywords = await self._extract_keywords(document.extracted_text)
			document.ai_tags = keywords[:10]  # Top 10 keywords
			
			# Generate summary
			summary = await self._generate_summary(document.extracted_text)
			document.content_summary = summary
			
			# Calculate readability metrics
			document.business_metadata['readability_score'] = textstat.flesch_reading_ease(document.extracted_text)
			document.business_metadata['word_count'] = len(document.extracted_text.split())
			
		except Exception as e:
			print(f"Failed to analyze document content: {e}")
	
	async def _extract_keywords(self, text: str, max_keywords: int = 20) -> List[str]:
		"""Extract keywords using TF-IDF"""
		
		try:
			# Fit TF-IDF vectorizer
			tfidf_matrix = self.vectorizer.fit_transform([text])
			feature_names = self.vectorizer.get_feature_names_out()
			
			# Get scores
			scores = tfidf_matrix.toarray()[0]
			
			# Get top keywords
			keyword_scores = list(zip(feature_names, scores))
			keyword_scores.sort(key=lambda x: x[1], reverse=True)
			
			return [keyword for keyword, score in keyword_scores[:max_keywords] if score > 0]
		
		except Exception as e:
			print(f"Failed to extract keywords: {e}")
			return []
	
	async def _generate_summary(self, text: str, max_sentences: int = 3) -> str:
		"""Generate document summary using extractive summarization"""
		
		try:
			# Simple extractive summarization
			sentences = text.split('.')
			if len(sentences) <= max_sentences:
				return text
			
			# Score sentences by keyword density
			keywords = await self._extract_keywords(text, 10)
			sentence_scores = []
			
			for i, sentence in enumerate(sentences):
				score = sum(1 for keyword in keywords if keyword.lower() in sentence.lower())
				sentence_scores.append((i, score, sentence.strip()))
			
			# Get top sentences
			sentence_scores.sort(key=lambda x: x[1], reverse=True)
			top_sentences = sorted(sentence_scores[:max_sentences], key=lambda x: x[0])
			
			return '. '.join([sentence for _, _, sentence in top_sentences if sentence])
		
		except Exception as e:
			print(f"Failed to generate summary: {e}")
			return text[:500] + "..." if len(text) > 500 else text
	
	async def _create_search_index(self, document: DCMDocument):
		"""Create search index entry for document"""
		
		search_index = DCMSearchIndex(
			tenant_id=document.tenant_id,
			created_by=document.created_by,
			updated_by=document.updated_by,
			document_id=document.id,
			content_hash=document.file_hash,
			extracted_text=document.extracted_text or "",
			metadata_text=f"{document.name} {document.title} {document.description or ''} {' '.join(document.keywords)}",
			title=document.title,
			tags=document.keywords + document.ai_tags,
			language=document.language
		)
		
		# Generate search-optimized keywords
		if document.extracted_text:
			search_index.search_keywords = await self._extract_keywords(document.extracted_text, 50)
		
		# Store in search index (mock)
		await self._store_search_index(search_index)
	
	async def _check_permission(self, resource_id: str, user_id: str, required_level: DCMPermissionLevel) -> bool:
		"""Check if user has required permission level"""
		
		# Check cache first
		cache_key = f"{resource_id}_{user_id}"
		if cache_key in self._permission_cache:
			cached_level = self._permission_cache[cache_key]
			return self._permission_level_sufficient(cached_level, required_level)
		
		# Mock permission check - in real implementation, query database
		# For now, assume all users have READ access, creators have ADMIN access
		permission_level = DCMPermissionLevel.READ
		
		# Cache result
		self._permission_cache[cache_key] = permission_level
		
		return self._permission_level_sufficient(permission_level, required_level)
	
	def _permission_level_sufficient(self, user_level: DCMPermissionLevel, required_level: DCMPermissionLevel) -> bool:
		"""Check if user permission level is sufficient"""
		
		level_hierarchy = {
			DCMPermissionLevel.NONE: 0,
			DCMPermissionLevel.READ: 1,
			DCMPermissionLevel.WRITE: 2,
			DCMPermissionLevel.DELETE: 3,
			DCMPermissionLevel.ADMIN: 4,
			DCMPermissionLevel.OWNER: 5
		}
		
		return level_hierarchy.get(user_level, 0) >= level_hierarchy.get(required_level, 0)
	
	def _calculate_permission_flags(self, permission_level: DCMPermissionLevel) -> Dict[str, bool]:
		"""Calculate individual permission flags based on level"""
		
		flags = {
			'can_read': False,
			'can_write': False,
			'can_delete': False,
			'can_share': False,
			'can_download': False,
			'can_print': False,
			'can_export': False,
			'can_comment': False,
			'can_approve': False
		}
		
		if permission_level in [DCMPermissionLevel.READ, DCMPermissionLevel.WRITE, 
								DCMPermissionLevel.DELETE, DCMPermissionLevel.ADMIN, DCMPermissionLevel.OWNER]:
			flags.update({
				'can_read': True,
				'can_download': True,
				'can_print': True,
				'can_comment': True
			})
		
		if permission_level in [DCMPermissionLevel.WRITE, DCMPermissionLevel.DELETE, 
								DCMPermissionLevel.ADMIN, DCMPermissionLevel.OWNER]:
			flags.update({
				'can_write': True,
				'can_share': True,
				'can_export': True
			})
		
		if permission_level in [DCMPermissionLevel.DELETE, DCMPermissionLevel.ADMIN, DCMPermissionLevel.OWNER]:
			flags['can_delete'] = True
		
		if permission_level in [DCMPermissionLevel.ADMIN, DCMPermissionLevel.OWNER]:
			flags['can_approve'] = True
		
		return flags
	
	async def _log_audit_event(self, **kwargs):
		"""Log audit event"""
		
		audit_log = DCMAuditLog(
			tenant_id=kwargs['tenant_id'],
			created_by=kwargs['user_id'],
			updated_by=kwargs['user_id'],
			**kwargs
		)
		
		# Store audit log (mock)
		print(f"Audit: {audit_log.action} on {audit_log.resource_type} {audit_log.resource_id} by user {audit_log.user_id}")
	
	async def _send_notification(self, **kwargs):
		"""Send notification to user"""
		
		notification = DCMNotification(
			tenant_id=kwargs['tenant_id'],
			created_by='system',
			updated_by='system',
			**kwargs
		)
		
		# Send notification (mock)
		print(f"Notification: {notification.title} to user {notification.recipient_id}")
	
	# Mock database methods (in real implementation, these would interact with actual database)
	async def _get_document_from_db(self, document_id: str, tenant_id: str) -> Optional[DCMDocument]:
		"""Mock method to get document from database"""
		return None  # Would return actual document from database
	
	async def _get_document_versions_from_db(self, document_id: str, tenant_id: str) -> List[DCMDocumentVersion]:
		"""Mock method to get document versions from database"""
		return []
	
	async def _get_workflow_from_db(self, workflow_id: str, tenant_id: str) -> Optional[DCMWorkflow]:
		"""Mock method to get workflow from database"""
		return None
	
	async def _store_search_index(self, search_index: DCMSearchIndex):
		"""Mock method to store search index"""
		pass
	
	# Additional helper methods would be implemented here for a complete service
	def _extract_mentions(self, text: str) -> List[str]:
		"""Extract @mentions from text"""
		import re
		mentions = re.findall(r'@(\w+)', text)
		return mentions
	
	def _parse_search_query(self, query: str) -> List[str]:
		"""Parse search query into terms"""
		return query.lower().split()
	
	def _build_search_filters(self, filters: Dict[str, Any], tenant_id: str) -> Dict[str, Any]:
		"""Build search filters"""
		filters['tenant_id'] = tenant_id
		return filters
	
	async def _execute_document_search(self, **kwargs) -> Dict[str, Any]:
		"""Execute document search"""
		return {'documents': [], 'search_time_ms': 0}
	
	def _calculate_search_relevance(self, results: List[Dict], terms: List[str]) -> List[Dict]:
		"""Calculate search relevance scores"""
		return results
	
	async def _generate_search_suggestions(self, query: str, tenant_id: str) -> List[str]:
		"""Generate search suggestions"""
		return []