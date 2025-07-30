"""
APG Document Content Management - Intelligent Document Processing Service

Self-learning AI-powered document processing with continuous improvement,
OCR, data extraction, and automated classification.

Copyright © 2025 Datacraft
Author: Nyimbi Odero <nyimbi@gmail.com>
Website: www.datacraft.co.ke
"""

import asyncio
import hashlib
import json
import logging
import mimetypes
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union
from uuid import UUID

import aiofiles
from PIL import Image
import pytesseract
import fitz  # PyMuPDF
import pandas as pd
from docx import Document as DocxDocument

from .models import (
	DCMDocument, DCMIntelligentProcessing, DCMContentIntelligence,
	DCMContentFormat, DCMDocumentType, ValidatedConfidenceScore
)


class IDPProcessor:
	"""Intelligent Document Processing with self-learning AI capabilities"""
	
	def __init__(self, apg_ai_client=None, apg_rag_client=None):
		"""Initialize IDP processor with APG AI/ML integration"""
		self.apg_ai_client = apg_ai_client
		self.apg_rag_client = apg_rag_client
		self.logger = logging.getLogger(__name__)
		
		# Supported formats for processing
		self.ocr_formats = {'.pdf', '.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif'}
		self.text_formats = {'.txt', '.md', '.csv', '.json', '.xml', '.html'}
		self.office_formats = {'.docx', '.xlsx', '.pptx'}
		
		# Model performance tracking
		self.model_stats = {
			'processing_count': 0,
			'accuracy_scores': [],
			'processing_times': [],
			'error_count': 0
		}
	
	async def process_document(
		self, 
		document: DCMDocument,
		file_path: str,
		processing_options: Optional[Dict[str, Any]] = None
	) -> DCMIntelligentProcessing:
		"""Process document with intelligent AI-powered extraction"""
		start_time = time.time()
		
		try:
			# Determine processing strategy based on content format
			processing_type = self._determine_processing_type(document.content_format)
			
			# Extract content based on format
			if processing_type == "ocr":
				extracted_data = await self._process_with_ocr(file_path, document)
			elif processing_type == "extraction":
				extracted_data = await self._process_structured_document(file_path, document)
			else:
				extracted_data = await self._process_text_document(file_path, document)
			
			# Apply AI classification and enhancement
			classification_results = await self._classify_content(extracted_data, document)
			
			# Create processing record
			processing_record = DCMIntelligentProcessing(
				tenant_id=document.tenant_id,
				created_by=document.created_by,
				updated_by=document.updated_by,
				document_id=document.id,
				processing_type=processing_type,
				model_version=await self._get_current_model_version(),
				ai_model_id=await self._get_ai_model_id(),
				confidence_score=self._calculate_confidence_score(extracted_data, classification_results),
				extracted_data=extracted_data,
				classification_results=classification_results,
				validation_status="pending",
				user_corrections=[],
				processing_time_ms=int((time.time() - start_time) * 1000),
				accuracy_score=None,
				error_count=0
			)
			
			# Update model statistics
			self._update_model_stats(processing_record)
			
			self.logger.info(f"Successfully processed document {document.id} in {processing_record.processing_time_ms}ms")
			return processing_record
			
		except Exception as e:
			self.logger.error(f"Error processing document {document.id}: {str(e)}")
			
			# Create error processing record
			return DCMIntelligentProcessing(
				tenant_id=document.tenant_id,
				created_by=document.created_by,
				updated_by=document.updated_by,
				document_id=document.id,
				processing_type="error",
				model_version=await self._get_current_model_version(),
				ai_model_id=await self._get_ai_model_id(),
				confidence_score=0.0,
				extracted_data={"error": str(e)},
				classification_results=[],
				validation_status="failed",
				processing_time_ms=int((time.time() - start_time) * 1000),
				error_count=1
			)
	
	async def _process_with_ocr(self, file_path: str, document: DCMDocument) -> Dict[str, Any]:
		"""Process document using OCR for text extraction"""
		extracted_data = {
			"text_content": "",
			"pages": [],
			"images": [],
			"tables": [],
			"confidence_scores": []
		}
		
		try:
			if document.content_format == DCMContentFormat.PDF:
				# Use PyMuPDF for PDF processing
				doc = fitz.open(file_path)
				
				for page_num in range(doc.page_count):
					page = doc[page_num]
					
					# Extract text directly from PDF
					text = page.get_text()
					
					# If no text found, use OCR on page image
					if not text.strip():
						pix = page.get_pixmap()
						img_data = pix.tobytes("ppm")
						text = pytesseract.image_to_string(Image.open(io.BytesIO(img_data)))
					
					# Extract tables
					tables = page.find_tables()
					page_tables = []
					for table in tables:
						table_data = table.extract()
						page_tables.append({
							"data": table_data,
							"bbox": table.bbox
						})
					
					# Extract images
					image_list = page.get_images()
					page_images = []
					for img_index, img in enumerate(image_list):
						xref = img[0]
						base_image = doc.extract_image(xref)
						page_images.append({
							"index": img_index,
							"ext": base_image["ext"],
							"width": base_image["width"],
							"height": base_image["height"]
						})
					
					extracted_data["pages"].append({
						"page_number": page_num + 1,
						"text": text,
						"word_count": len(text.split()),
						"tables": page_tables,
						"images": page_images
					})
					
					extracted_data["text_content"] += text + "\n"
				
				doc.close()
				
			else:
				# Process image files with OCR
				with Image.open(file_path) as img:
					# Use Tesseract for OCR
					text = pytesseract.image_to_string(img)
					confidence_data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
					
					extracted_data["text_content"] = text
					extracted_data["pages"].append({
						"page_number": 1,
						"text": text,
						"word_count": len(text.split()),
						"ocr_confidence": confidence_data
					})
			
			# Apply APG AI post-processing if available
			if self.apg_ai_client:
				enhanced_data = await self._apply_ai_postprocessing(extracted_data)
				extracted_data.update(enhanced_data)
				
		except Exception as e:
			self.logger.error(f"OCR processing error: {str(e)}")
			extracted_data["error"] = str(e)
		
		return extracted_data
	
	async def _process_structured_document(self, file_path: str, document: DCMDocument) -> Dict[str, Any]:
		"""Process structured documents (Office formats)"""
		extracted_data = {
			"text_content": "",
			"metadata": {},
			"structure": {},
			"entities": []
		}
		
		try:
			if document.content_format == DCMContentFormat.DOCX:
				# Process Word document
				doc = DocxDocument(file_path)
				
				# Extract text content
				full_text = []
				for paragraph in doc.paragraphs:
					full_text.append(paragraph.text)
				extracted_data["text_content"] = "\n".join(full_text)
				
				# Extract document properties
				props = doc.core_properties
				extracted_data["metadata"] = {
					"title": props.title or "",
					"author": props.author or "",
					"subject": props.subject or "",
					"created": props.created.isoformat() if props.created else None,
					"modified": props.modified.isoformat() if props.modified else None,
					"word_count": len(extracted_data["text_content"].split())
				}
				
				# Extract structure (headings, styles)
				structure = {
					"headings": [],
					"paragraphs": len(doc.paragraphs),
					"styles": []
				}
				
				for paragraph in doc.paragraphs:
					if paragraph.style.name.startswith('Heading'):
						structure["headings"].append({
							"level": paragraph.style.name,
							"text": paragraph.text
						})
				
				extracted_data["structure"] = structure
				
			elif document.content_format == DCMContentFormat.XLSX:
				# Process Excel document
				df = pd.read_excel(file_path, sheet_name=None)
				
				extracted_data["structure"] = {
					"sheets": list(df.keys()),
					"total_rows": sum(sheet.shape[0] for sheet in df.values()),
					"total_columns": sum(sheet.shape[1] for sheet in df.values())
				}
				
				# Convert to text for analysis
				text_parts = []
				for sheet_name, sheet_data in df.items():
					text_parts.append(f"Sheet: {sheet_name}")
					text_parts.append(sheet_data.to_string())
				
				extracted_data["text_content"] = "\n".join(text_parts)
			
			# Apply APG AI entity extraction if available
			if self.apg_ai_client:
				entities = await self._extract_entities(extracted_data["text_content"])
				extracted_data["entities"] = entities
				
		except Exception as e:
			self.logger.error(f"Structured document processing error: {str(e)}")
			extracted_data["error"] = str(e)
		
		return extracted_data
	
	async def _process_text_document(self, file_path: str, document: DCMDocument) -> Dict[str, Any]:
		"""Process plain text documents"""
		extracted_data = {
			"text_content": "",
			"metadata": {},
			"statistics": {}
		}
		
		try:
			async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
				content = await f.read()
			
			extracted_data["text_content"] = content
			
			# Calculate text statistics
			words = content.split()
			lines = content.split('\n')
			
			extracted_data["statistics"] = {
				"character_count": len(content),
				"word_count": len(words),
				"line_count": len(lines),
				"paragraph_count": len([p for p in content.split('\n\n') if p.strip()]),
				"average_words_per_line": len(words) / len(lines) if lines else 0
			}
			
			# Detect language if APG AI available
			if self.apg_ai_client:
				language = await self._detect_language(content)
				extracted_data["metadata"]["detected_language"] = language
				
		except Exception as e:
			self.logger.error(f"Text document processing error: {str(e)}")
			extracted_data["error"] = str(e)
		
		return extracted_data
	
	async def _classify_content(self, extracted_data: Dict[str, Any], document: DCMDocument) -> List[Dict[str, Any]]:
		"""Apply AI-powered content classification"""
		classification_results = []
		
		try:
			if not extracted_data.get("text_content"):
				return classification_results
			
			text_content = extracted_data["text_content"]
			
			# Document type classification
			if self.apg_ai_client:
				doc_type_result = await self._classify_document_type(text_content)
				classification_results.append({
					"type": "document_type",
					"classification": doc_type_result["type"],
					"confidence": doc_type_result["confidence"],
					"alternatives": doc_type_result.get("alternatives", [])
				})
			
			# Content category classification
			if self.apg_ai_client:
				category_result = await self._classify_content_category(text_content)
				classification_results.append({
					"type": "content_category",
					"classification": category_result["category"],
					"confidence": category_result["confidence"],
					"subcategories": category_result.get("subcategories", [])
				})
			
			# Sensitivity classification
			sensitivity_result = await self._classify_sensitivity(text_content)
			classification_results.append({
				"type": "sensitivity",
				"classification": sensitivity_result["level"],
				"confidence": sensitivity_result["confidence"],
				"detected_patterns": sensitivity_result.get("patterns", [])
			})
			
			# Industry-specific classification
			if self.apg_ai_client:
				industry_result = await self._classify_industry_context(text_content)
				classification_results.append({
					"type": "industry_context",
					"classification": industry_result["industry"],
					"confidence": industry_result["confidence"],
					"specialized_terms": industry_result.get("terms", [])
				})
				
		except Exception as e:
			self.logger.error(f"Content classification error: {str(e)}")
			classification_results.append({
				"type": "error",
				"error": str(e)
			})
		
		return classification_results
	
	async def _apply_ai_postprocessing(self, extracted_data: Dict[str, Any]) -> Dict[str, Any]:
		"""Apply APG AI post-processing to improve extraction quality"""
		enhanced_data = {}
		
		if self.apg_ai_client and extracted_data.get("text_content"):
			try:
				# Spell check and correction
				corrected_text = await self.apg_ai_client.spell_check(
					extracted_data["text_content"]
				)
				enhanced_data["corrected_text"] = corrected_text
				
				# Text enhancement and cleanup
				cleaned_text = await self.apg_ai_client.clean_text(
					extracted_data["text_content"]
				)
				enhanced_data["cleaned_text"] = cleaned_text
				
				# Structure detection
				structure = await self.apg_ai_client.detect_structure(
					extracted_data["text_content"]
				)
				enhanced_data["ai_detected_structure"] = structure
				
			except Exception as e:
				self.logger.error(f"AI post-processing error: {str(e)}")
		
		return enhanced_data
	
	async def _extract_entities(self, text: str) -> List[Dict[str, Any]]:
		"""Extract named entities using APG AI"""
		entities = []
		
		if self.apg_ai_client:
			try:
				# Use APG AI for entity extraction
				extracted_entities = await self.apg_ai_client.extract_entities(text)
				
				for entity in extracted_entities:
					entities.append({
						"text": entity["text"],
						"label": entity["label"],
						"confidence": entity["confidence"],
						"start": entity.get("start", 0),
						"end": entity.get("end", 0)
					})
					
			except Exception as e:
				self.logger.error(f"Entity extraction error: {str(e)}")
		
		return entities
	
	async def _detect_language(self, text: str) -> str:
		"""Detect document language using APG AI"""
		try:
			if self.apg_ai_client:
				result = await self.apg_ai_client.detect_language(text)
				return result.get("language", "en")
		except Exception as e:
			self.logger.error(f"Language detection error: {str(e)}")
		
		return "en"  # Default to English
	
	async def _classify_document_type(self, text: str) -> Dict[str, Any]:
		"""Classify document type using APG AI"""
		try:
			if self.apg_ai_client:
				result = await self.apg_ai_client.classify_document_type(text)
				return {
					"type": result.get("type", "unknown"),
					"confidence": result.get("confidence", 0.5),
					"alternatives": result.get("alternatives", [])
				}
		except Exception as e:
			self.logger.error(f"Document type classification error: {str(e)}")
		
		return {"type": "unknown", "confidence": 0.0, "alternatives": []}
	
	async def _classify_content_category(self, text: str) -> Dict[str, Any]:
		"""Classify content category using APG AI"""
		try:
			if self.apg_ai_client:
				result = await self.apg_ai_client.classify_content_category(text)
				return {
					"category": result.get("category", "general"),
					"confidence": result.get("confidence", 0.5),
					"subcategories": result.get("subcategories", [])
				}
		except Exception as e:
			self.logger.error(f"Content category classification error: {str(e)}")
		
		return {"category": "general", "confidence": 0.0, "subcategories": []}
	
	async def _classify_sensitivity(self, text: str) -> Dict[str, Any]:
		"""Classify content sensitivity level"""
		patterns = {
			"confidential": ["confidential", "proprietary", "internal only", "restricted"],
			"personal": ["social security", "ssn", "credit card", "phone number", "email"],
			"financial": ["account number", "routing number", "tax id", "invoice"],
			"medical": ["patient", "diagnosis", "medical record", "hipaa"]
		}
		
		text_lower = text.lower()
		detected_patterns = []
		sensitivity_level = "public"
		confidence = 0.6
		
		for level, keywords in patterns.items():
			for keyword in keywords:
				if keyword in text_lower:
					detected_patterns.append(keyword)
					if level in ["confidential", "personal"]:
						sensitivity_level = "confidential"
						confidence = 0.8
					elif level in ["financial", "medical"] and sensitivity_level == "public":
						sensitivity_level = "restricted"
						confidence = 0.7
		
		return {
			"level": sensitivity_level,
			"confidence": confidence,
			"patterns": detected_patterns
		}
	
	async def _classify_industry_context(self, text: str) -> Dict[str, Any]:
		"""Classify industry context using APG AI"""
		try:
			if self.apg_ai_client:
				result = await self.apg_ai_client.classify_industry(text)
				return {
					"industry": result.get("industry", "general"),
					"confidence": result.get("confidence", 0.5),
					"terms": result.get("specialized_terms", [])
				}
		except Exception as e:
			self.logger.error(f"Industry classification error: {str(e)}")
		
		return {"industry": "general", "confidence": 0.0, "terms": []}
	
	def _determine_processing_type(self, content_format: DCMContentFormat) -> str:
		"""Determine the appropriate processing type for the content format"""
		format_str = content_format.value.lower()
		
		if format_str in {'pdf', 'png', 'jpg', 'jpeg', 'tiff', 'bmp', 'gif'}:
			return "ocr"
		elif format_str in {'docx', 'xlsx', 'pptx'}:
			return "extraction"
		else:
			return "text"
	
	def _calculate_confidence_score(
		self, 
		extracted_data: Dict[str, Any], 
		classification_results: List[Dict[str, Any]]
	) -> float:
		"""Calculate overall confidence score for processing results"""
		confidence_scores = []
		
		# Base confidence on successful extraction
		if extracted_data.get("text_content"):
			confidence_scores.append(0.8)
		
		# Include classification confidences
		for result in classification_results:
			if "confidence" in result:
				confidence_scores.append(result["confidence"])
		
		# OCR confidence if available
		if "confidence_scores" in extracted_data:
			ocr_scores = extracted_data["confidence_scores"]
			if ocr_scores:
				avg_ocr_confidence = sum(ocr_scores) / len(ocr_scores) / 100.0
				confidence_scores.append(avg_ocr_confidence)
		
		return sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0.0
	
	async def _get_current_model_version(self) -> str:
		"""Get current AI model version"""
		if self.apg_ai_client:
			try:
				return await self.apg_ai_client.get_model_version()
			except:
				pass
		return "1.0.0"
	
	async def _get_ai_model_id(self) -> str:
		"""Get APG AI model ID"""
		if self.apg_ai_client:
			try:
				return await self.apg_ai_client.get_model_id()
			except:
				pass
		return "apg-idp-base-v1"
	
	def _update_model_stats(self, processing_record: DCMIntelligentProcessing):
		"""Update model performance statistics"""
		self.model_stats['processing_count'] += 1
		self.model_stats['processing_times'].append(processing_record.processing_time_ms)
		
		if processing_record.accuracy_score:
			self.model_stats['accuracy_scores'].append(processing_record.accuracy_score)
		
		if processing_record.error_count > 0:
			self.model_stats['error_count'] += 1
	
	async def apply_user_corrections(
		self, 
		processing_id: str, 
		corrections: List[Dict[str, Any]]
	) -> DCMIntelligentProcessing:
		"""Apply user corrections for self-learning improvement"""
		# This would update the processing record with corrections
		# and potentially trigger model retraining
		pass
	
	async def get_processing_analytics(self) -> Dict[str, Any]:
		"""Get processing performance analytics"""
		return {
			"total_processed": self.model_stats['processing_count'],
			"average_processing_time": (
				sum(self.model_stats['processing_times']) / 
				len(self.model_stats['processing_times'])
			) if self.model_stats['processing_times'] else 0,
			"average_accuracy": (
				sum(self.model_stats['accuracy_scores']) / 
				len(self.model_stats['accuracy_scores'])
			) if self.model_stats['accuracy_scores'] else 0,
			"error_rate": (
				self.model_stats['error_count'] / 
				self.model_stats['processing_count']
			) if self.model_stats['processing_count'] > 0 else 0
		}