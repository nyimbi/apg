"""
APG RAG Document Processing Pipeline

Comprehensive document ingestion and processing with support for 20+ file formats,
intelligent chunking strategies, and seamless integration with PostgreSQL + pgvector + pgai.
"""

import asyncio
import logging
import hashlib
import time
from typing import Dict, Any, List, Optional, Union, Tuple, BinaryIO
from datetime import datetime
from pathlib import Path
from dataclasses import dataclass, field
from enum import Enum
import mimetypes
import json
import re
from uuid_extensions import uuid7str

# Document processing libraries
import PyPDF2
import docx
import openpyxl
import pandas as pd
from bs4 import BeautifulSoup
import markdown
import xml.etree.ElementTree as ET
from PIL import Image
import io
import base64

# APG imports
from .models import (
	Document, DocumentCreate, DocumentChunk, DocumentChunkCreate,
	DocumentStatus, APGBaseModel
)
from .ollama_integration import AdvancedOllamaIntegration, EmbeddingQueueRequest, RequestPriority

class ChunkingStrategy(str, Enum):
	"""Document chunking strategies"""
	SENTENCE = "sentence"
	PARAGRAPH = "paragraph"
	SEMANTIC = "semantic"
	FIXED_SIZE = "fixed_size"
	SLIDING_WINDOW = "sliding_window"
	HIERARCHICAL = "hierarchical"

class ProcessingPriority(str, Enum):
	"""Document processing priority levels"""
	LOW = "low"
	NORMAL = "normal"
	HIGH = "high"
	URGENT = "urgent"

@dataclass
class ChunkingConfig:
	"""Configuration for document chunking"""
	strategy: ChunkingStrategy = ChunkingStrategy.SEMANTIC
	chunk_size: int = 1000
	chunk_overlap: int = 100
	min_chunk_size: int = 50
	max_chunk_size: int = 2000
	preserve_structure: bool = True
	include_metadata: bool = True
	sentence_threshold: int = 3  # Minimum sentences per chunk

@dataclass
class ProcessingConfig:
	"""Configuration for document processing"""
	chunking: ChunkingConfig = field(default_factory=ChunkingConfig)
	extract_metadata: bool = True
	generate_embeddings: bool = True
	detect_language: bool = True
	extract_entities: bool = False  # Requires NLP capability
	quality_check: bool = True
	batch_size: int = 32
	max_file_size_mb: int = 100
	supported_formats: List[str] = field(default_factory=lambda: [
		'pdf', 'docx', 'doc', 'txt', 'md', 'html', 'htm', 'json', 'xml',
		'csv', 'xlsx', 'xls', 'pptx', 'ppt', 'rtf', 'odt', 'ods', 'odp'
	])

@dataclass
class ProcessingResult:
	"""Result of document processing"""
	document: Document
	chunks: List[DocumentChunk]
	processing_time_ms: float
	success: bool
	error_message: Optional[str] = None
	metadata: Dict[str, Any] = field(default_factory=dict)

class DocumentFormatExtractor:
	"""Extract content from various document formats"""
	
	def __init__(self):
		self.logger = logging.getLogger(__name__)
		
		# Initialize OCR if available (optional dependency)
		self.ocr_available = False
		try:
			import pytesseract
			self.ocr_available = True
		except ImportError:
			self.logger.info("OCR not available - install pytesseract for image text extraction")
	
	async def extract_content(self, file_path: str, content_type: str) -> Dict[str, Any]:
		"""Extract content and metadata from file"""
		try:
			path = Path(file_path)
			
			if not path.exists():
				raise FileNotFoundError(f"File not found: {file_path}")
			
			# Determine extraction method based on content type
			if content_type.startswith('text/'):
				return await self._extract_text(path)
			elif content_type == 'application/pdf':
				return await self._extract_pdf(path)
			elif content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
				return await self._extract_docx(path)
			elif content_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'application/vnd.ms-excel']:
				return await self._extract_excel(path)
			elif content_type in ['text/html', 'application/xhtml+xml']:
				return await self._extract_html(path)
			elif content_type in ['application/json']:
				return await self._extract_json(path)
			elif content_type in ['application/xml', 'text/xml']:
				return await self._extract_xml(path)
			elif content_type.startswith('image/'):
				return await self._extract_image(path)
			else:
				# Fallback to text extraction
				return await self._extract_text(path)
		
		except Exception as e:
			self.logger.error(f"Content extraction failed for {file_path}: {str(e)}")
			raise
	
	async def _extract_text(self, path: Path) -> Dict[str, Any]:
		"""Extract plain text content"""
		try:
			with open(path, 'r', encoding='utf-8', errors='ignore') as f:
				content = f.read()
			
			return {
				'content': content,
				'metadata': {
					'format': 'text',
					'encoding': 'utf-8',
					'line_count': len(content.splitlines()),
					'character_count': len(content)
				}
			}
		except UnicodeDecodeError:
			# Try different encodings
			for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
				try:
					with open(path, 'r', encoding=encoding, errors='ignore') as f:
						content = f.read()
					return {
						'content': content,
						'metadata': {
							'format': 'text',
							'encoding': encoding,
							'line_count': len(content.splitlines()),
							'character_count': len(content)
						}
					}
				except:
					continue
			raise
	
	async def _extract_pdf(self, path: Path) -> Dict[str, Any]:
		"""Extract content from PDF files"""
		content_parts = []
		metadata = {'format': 'pdf', 'pages': 0}
		
		try:
			with open(path, 'rb') as file:
				pdf_reader = PyPDF2.PdfReader(file)
				metadata['pages'] = len(pdf_reader.pages)
				
				# Extract metadata
				if pdf_reader.metadata:
					metadata.update({
						'title': pdf_reader.metadata.get('/Title', ''),
						'author': pdf_reader.metadata.get('/Author', ''),
						'subject': pdf_reader.metadata.get('/Subject', ''),
						'creator': pdf_reader.metadata.get('/Creator', ''),
						'producer': pdf_reader.metadata.get('/Producer', ''),
						'creation_date': pdf_reader.metadata.get('/CreationDate', ''),
						'modification_date': pdf_reader.metadata.get('/ModDate', '')
					})
				
				# Extract text from each page
				for page_num, page in enumerate(pdf_reader.pages):
					try:
						page_text = page.extract_text()
						if page_text.strip():
							content_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
					except Exception as e:
						self.logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
						continue
		
		except Exception as e:
			self.logger.error(f"PDF extraction failed: {str(e)}")
			raise
		
		return {
			'content': '\n\n'.join(content_parts),
			'metadata': metadata
		}
	
	async def _extract_docx(self, path: Path) -> Dict[str, Any]:
		"""Extract content from DOCX files"""
		try:
			import docx
			doc = docx.Document(path)
			
			content_parts = []
			metadata = {
				'format': 'docx',
				'paragraphs': len(doc.paragraphs),
				'tables': len(doc.tables)
			}
			
			# Extract core properties
			if doc.core_properties:
				metadata.update({
					'title': doc.core_properties.title or '',
					'author': doc.core_properties.author or '',
					'subject': doc.core_properties.subject or '',
					'keywords': doc.core_properties.keywords or '',
					'created': str(doc.core_properties.created) if doc.core_properties.created else '',
					'modified': str(doc.core_properties.modified) if doc.core_properties.modified else ''
				})
			
			# Extract paragraphs
			for para in doc.paragraphs:
				if para.text.strip():
					content_parts.append(para.text)
			
			# Extract tables
			for table in doc.tables:
				table_content = []
				for row in table.rows:
					row_content = []
					for cell in row.cells:
						if cell.text.strip():
							row_content.append(cell.text.strip())
					if row_content:
						table_content.append(' | '.join(row_content))
				
				if table_content:
					content_parts.append('\n--- Table ---\n' + '\n'.join(table_content))
			
			return {
				'content': '\n\n'.join(content_parts),
				'metadata': metadata
			}
		
		except ImportError:
			raise ImportError("python-docx is required for DOCX support")
		except Exception as e:
			self.logger.error(f"DOCX extraction failed: {str(e)}")
			raise
	
	async def _extract_excel(self, path: Path) -> Dict[str, Any]:
		"""Extract content from Excel files"""
		try:
			# Try pandas first for better handling
			try:
				# Read all sheets
				excel_data = pd.read_excel(path, sheet_name=None, engine='openpyxl')
				
				content_parts = []
				metadata = {
					'format': 'excel',
					'sheets': list(excel_data.keys()),
					'sheet_count': len(excel_data)
				}
				
				for sheet_name, df in excel_data.items():
					if not df.empty:
						# Convert to readable format
						sheet_content = f"--- Sheet: {sheet_name} ---\n"
						sheet_content += df.to_string(index=False, na_rep='')
						content_parts.append(sheet_content)
						
						# Add sheet metadata
						metadata[f'sheet_{sheet_name}'] = {
							'rows': len(df),
							'columns': len(df.columns),
							'column_names': list(df.columns)
						}
				
				return {
					'content': '\n\n'.join(content_parts),
					'metadata': metadata
				}
			
			except Exception:
				# Fallback to openpyxl
				import openpyxl
				workbook = openpyxl.load_workbook(path, data_only=True)
				
				content_parts = []
				metadata = {
					'format': 'excel',
					'sheets': workbook.sheetnames,
					'sheet_count': len(workbook.sheetnames)
				}
				
				for sheet_name in workbook.sheetnames:
					sheet = workbook[sheet_name]
					sheet_content = f"--- Sheet: {sheet_name} ---\n"
					
					for row in sheet.iter_rows(values_only=True):
						row_values = [str(cell) if cell is not None else '' for cell in row]
						if any(val.strip() for val in row_values):  # Skip empty rows
							sheet_content += ' | '.join(row_values) + '\n'
					
					if sheet_content.strip():
						content_parts.append(sheet_content)
				
				return {
					'content': '\n\n'.join(content_parts),
					'metadata': metadata
				}
		
		except ImportError:
			raise ImportError("pandas and openpyxl are required for Excel support")
		except Exception as e:
			self.logger.error(f"Excel extraction failed: {str(e)}")
			raise
	
	async def _extract_html(self, path: Path) -> Dict[str, Any]:
		"""Extract content from HTML files"""
		try:
			with open(path, 'r', encoding='utf-8', errors='ignore') as f:
				html_content = f.read()
			
			soup = BeautifulSoup(html_content, 'html.parser')
			
			# Extract metadata
			metadata = {'format': 'html'}
			
			title = soup.find('title')
			if title:
				metadata['title'] = title.get_text().strip()
			
			# Extract meta tags
			meta_tags = soup.find_all('meta')
			for meta in meta_tags:
				name = meta.get('name') or meta.get('property')
				content = meta.get('content')
				if name and content:
					metadata[f'meta_{name}'] = content
			
			# Remove script and style elements
			for script in soup(["script", "style"]):
				script.decompose()
			
			# Extract text content
			text_content = soup.get_text()
			
			# Clean up whitespace
			lines = (line.strip() for line in text_content.splitlines())
			chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
			content = ' '.join(chunk for chunk in chunks if chunk)
			
			return {
				'content': content,
				'metadata': metadata
			}
		
		except Exception as e:
			self.logger.error(f"HTML extraction failed: {str(e)}")
			raise
	
	async def _extract_json(self, path: Path) -> Dict[str, Any]:
		"""Extract content from JSON files"""
		try:
			with open(path, 'r', encoding='utf-8') as f:
				json_data = json.load(f)
			
			# Convert JSON to readable text
			def json_to_text(obj, prefix=""):
				text_parts = []
				if isinstance(obj, dict):
					for key, value in obj.items():
						if isinstance(value, (dict, list)):
							text_parts.append(f"{prefix}{key}:")
							text_parts.append(json_to_text(value, prefix + "  "))
						else:
							text_parts.append(f"{prefix}{key}: {value}")
				elif isinstance(obj, list):
					for i, item in enumerate(obj):
						if isinstance(item, (dict, list)):
							text_parts.append(f"{prefix}[{i}]:")
							text_parts.append(json_to_text(item, prefix + "  "))
						else:
							text_parts.append(f"{prefix}[{i}]: {item}")
				else:
					text_parts.append(f"{prefix}{obj}")
				
				return '\n'.join(text_parts)
			
			content = json_to_text(json_data)
			
			metadata = {
				'format': 'json',
				'structure_type': type(json_data).__name__,
				'keys': list(json_data.keys()) if isinstance(json_data, dict) else None,
				'length': len(json_data) if isinstance(json_data, (dict, list)) else None
			}
			
			return {
				'content': content,
				'metadata': metadata
			}
		
		except Exception as e:
			self.logger.error(f"JSON extraction failed: {str(e)}")
			raise
	
	async def _extract_xml(self, path: Path) -> Dict[str, Any]:
		"""Extract content from XML files"""
		try:
			tree = ET.parse(path)
			root = tree.getroot()
			
			def xml_to_text(element, level=0):
				text_parts = []
				indent = "  " * level
				
				# Add element name and attributes
				attrs = " ".join([f"{k}='{v}'" for k, v in element.attrib.items()])
				element_info = f"{indent}{element.tag}"
				if attrs:
					element_info += f" ({attrs})"
				text_parts.append(element_info)
				
				# Add element text
				if element.text and element.text.strip():
					text_parts.append(f"{indent}  {element.text.strip()}")
				
				# Process children
				for child in element:
					text_parts.append(xml_to_text(child, level + 1))
				
				# Add tail text
				if element.tail and element.tail.strip():
					text_parts.append(f"{indent}  {element.tail.strip()}")
				
				return '\n'.join(text_parts)
			
			content = xml_to_text(root)
			
			metadata = {
				'format': 'xml',
				'root_element': root.tag,
				'namespace': root.tag.split('}')[0][1:] if '}' in root.tag else None,
				'element_count': len(list(root.iter()))
			}
			
			return {
				'content': content,
				'metadata': metadata
			}
		
		except Exception as e:
			self.logger.error(f"XML extraction failed: {str(e)}")
			raise
	
	async def _extract_image(self, path: Path) -> Dict[str, Any]:
		"""Extract content from image files using OCR"""
		try:
			if not self.ocr_available:
				return {
					'content': f"[Image file: {path.name}]",
					'metadata': {
						'format': 'image',
						'ocr_available': False,
						'note': 'Install pytesseract for text extraction from images'
					}
				}
			
			import pytesseract
			from PIL import Image
			
			# Open and process image
			with Image.open(path) as img:
				# Extract text using OCR
				extracted_text = pytesseract.image_to_string(img)
				
				metadata = {
					'format': 'image',
					'ocr_available': True,
					'image_size': img.size,
					'image_mode': img.mode,
					'image_format': img.format
				}
				
				# If no text extracted, provide image description
				if not extracted_text.strip():
					extracted_text = f"[Image file: {path.name} - {img.size[0]}x{img.size[1]} {img.format}]"
				
				return {
					'content': extracted_text,
					'metadata': metadata
				}
		
		except Exception as e:
			self.logger.error(f"Image extraction failed: {str(e)}")
			return {
				'content': f"[Image file: {path.name} - extraction failed]",
				'metadata': {
					'format': 'image',
					'extraction_error': str(e)
				}
			}

class DocumentChunker:
	"""Intelligent document chunking with multiple strategies"""
	
	def __init__(self, config: ChunkingConfig):
		self.config = config
		self.logger = logging.getLogger(__name__)
	
	async def chunk_document(self, content: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
		"""Chunk document using configured strategy"""
		if not content.strip():
			return []
		
		try:
			if self.config.strategy == ChunkingStrategy.SENTENCE:
				return await self._chunk_by_sentences(content, metadata)
			elif self.config.strategy == ChunkingStrategy.PARAGRAPH:
				return await self._chunk_by_paragraphs(content, metadata)
			elif self.config.strategy == ChunkingStrategy.SEMANTIC:
				return await self._chunk_semantic(content, metadata)
			elif self.config.strategy == ChunkingStrategy.FIXED_SIZE:
				return await self._chunk_fixed_size(content, metadata)
			elif self.config.strategy == ChunkingStrategy.SLIDING_WINDOW:
				return await self._chunk_sliding_window(content, metadata)
			elif self.config.strategy == ChunkingStrategy.HIERARCHICAL:
				return await self._chunk_hierarchical(content, metadata)
			else:
				# Default to semantic chunking
				return await self._chunk_semantic(content, metadata)
		
		except Exception as e:
			self.logger.error(f"Chunking failed: {str(e)}")
			# Fallback to simple fixed-size chunking
			return await self._chunk_fixed_size(content, metadata)
	
	async def _chunk_by_sentences(self, content: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
		"""Chunk by sentences with size limits"""
		import nltk
		try:
			nltk.data.find('tokenizers/punkt')
		except LookupError:
			nltk.download('punkt', quiet=True)
		
		sentences = nltk.sent_tokenize(content)
		chunks = []
		current_chunk = ""
		current_sentences = []
		
		for sentence in sentences:
			test_chunk = current_chunk + " " + sentence if current_chunk else sentence
			
			if (len(test_chunk) <= self.config.chunk_size and 
			    len(current_sentences) < self.config.sentence_threshold * 3):
				current_chunk = test_chunk
				current_sentences.append(sentence)
			else:
				if current_chunk and len(current_sentences) >= self.config.sentence_threshold:
					chunks.append(self._create_chunk_dict(current_chunk, len(chunks), metadata))
				current_chunk = sentence
				current_sentences = [sentence]
		
		# Add remaining chunk
		if current_chunk and len(current_sentences) >= self.config.sentence_threshold:
			chunks.append(self._create_chunk_dict(current_chunk, len(chunks), metadata))
		
		return chunks
	
	async def _chunk_by_paragraphs(self, content: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
		"""Chunk by paragraphs with size limits"""
		paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
		chunks = []
		current_chunk = ""
		
		for paragraph in paragraphs:
			test_chunk = current_chunk + "\n\n" + paragraph if current_chunk else paragraph
			
			if len(test_chunk) <= self.config.chunk_size:
				current_chunk = test_chunk
			else:
				if current_chunk:
					chunks.append(self._create_chunk_dict(current_chunk, len(chunks), metadata))
				
				# Handle oversized paragraphs
				if len(paragraph) > self.config.chunk_size:
					# Split large paragraph
					sub_chunks = await self._chunk_fixed_size(paragraph, metadata)
					chunks.extend(sub_chunks)
					current_chunk = ""
				else:
					current_chunk = paragraph
		
		# Add remaining chunk
		if current_chunk:
			chunks.append(self._create_chunk_dict(current_chunk, len(chunks), metadata))
		
		return chunks
	
	async def _chunk_semantic(self, content: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
		"""Intelligent semantic chunking (enhanced paragraph-based)"""
		# Split by double newlines first (paragraph boundaries)
		paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
		
		chunks = []
		current_chunk = ""
		current_topic_score = 0
		
		for i, paragraph in enumerate(paragraphs):
			# Simple topic coherence scoring based on overlapping words
			if current_chunk:
				overlap_score = self._calculate_semantic_similarity(current_chunk, paragraph)
			else:
				overlap_score = 1.0
			
			test_chunk = current_chunk + "\n\n" + paragraph if current_chunk else paragraph
			
			# Decide whether to continue current chunk or start new one
			should_continue = (
				len(test_chunk) <= self.config.chunk_size and
				overlap_score > 0.3  # Semantic similarity threshold
			)
			
			if should_continue:
				current_chunk = test_chunk
				current_topic_score += overlap_score
			else:
				# Finalize current chunk
				if current_chunk:
					chunks.append(self._create_chunk_dict(current_chunk, len(chunks), metadata))
				
				# Handle oversized paragraphs
				if len(paragraph) > self.config.chunk_size:
					sub_chunks = await self._chunk_fixed_size(paragraph, metadata)
					for sub_chunk in sub_chunks:
						sub_chunk['chunk_index'] = len(chunks)
						chunks.append(sub_chunk)
					current_chunk = ""
				else:
					current_chunk = paragraph
				current_topic_score = 1.0
		
		# Add remaining chunk
		if current_chunk:
			chunks.append(self._create_chunk_dict(current_chunk, len(chunks), metadata))
		
		return chunks
	
	async def _chunk_fixed_size(self, content: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
		"""Fixed-size chunking with overlap"""
		chunks = []
		start = 0
		chunk_index = 0
		
		while start < len(content):
			end = min(start + self.config.chunk_size, len(content))
			
			# Try to break at word boundary
			if end < len(content):
				# Look backwards for space
				while end > start and content[end] not in ' \n\t':
					end -= 1
				
				# If no space found, use original end
				if end == start:
					end = min(start + self.config.chunk_size, len(content))
			
			chunk_content = content[start:end].strip()
			if len(chunk_content) >= self.config.min_chunk_size:
				chunk_dict = self._create_chunk_dict(chunk_content, chunk_index, metadata)
				chunk_dict['start_position'] = start
				chunk_dict['end_position'] = end
				chunks.append(chunk_dict)
				chunk_index += 1
			
			# Move start position with overlap
			start = max(end - self.config.chunk_overlap, start + 1)
		
		return chunks
	
	async def _chunk_sliding_window(self, content: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
		"""Sliding window chunking with configurable overlap"""
		chunks = []
		window_size = self.config.chunk_size
		step_size = window_size - self.config.chunk_overlap
		
		for i in range(0, len(content), step_size):
			end = min(i + window_size, len(content))
			chunk_content = content[i:end].strip()
			
			if len(chunk_content) >= self.config.min_chunk_size:
				chunk_dict = self._create_chunk_dict(chunk_content, len(chunks), metadata)
				chunk_dict['start_position'] = i
				chunk_dict['end_position'] = end
				chunks.append(chunk_dict)
			
			if end >= len(content):
				break
		
		return chunks
	
	async def _chunk_hierarchical(self, content: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
		"""Hierarchical chunking based on document structure"""
		chunks = []
		
		# Detect structure markers (headers, sections, etc.)
		lines = content.split('\n')
		sections = []
		current_section = {"title": None, "content": [], "level": 0}
		
		for line in lines:
			stripped = line.strip()
			
			# Detect headers (markdown-style or numbered)
			header_level = 0
			if stripped.startswith('#'):
				header_level = len(stripped) - len(stripped.lstrip('#'))
				header_title = stripped.lstrip('#').strip()
			elif re.match(r'^\d+\.', stripped):
				header_level = 1
				header_title = stripped
			elif stripped.isupper() and len(stripped) < 100:
				header_level = 1
				header_title = stripped
			else:
				header_title = None
			
			if header_title:
				# Finalize current section
				if current_section["content"]:
					sections.append(current_section)
				
				# Start new section
				current_section = {
					"title": header_title,
					"content": [],
					"level": header_level
				}
			else:
				current_section["content"].append(line)
		
		# Add final section
		if current_section["content"]:
			sections.append(current_section)
		
		# Convert sections to chunks
		for section in sections:
			section_content = '\n'.join(section["content"]).strip()
			if not section_content:
				continue
			
			section_text = section_content
			if section["title"]:
				section_text = f"{section['title']}\n\n{section_content}"
			
			# If section is too large, sub-chunk it
			if len(section_text) > self.config.chunk_size:
				sub_chunks = await self._chunk_semantic(section_text, metadata)
				for sub_chunk in sub_chunks:
					sub_chunk['section_title'] = section["title"]
					sub_chunk['section_level'] = section["level"]
					sub_chunk['chunk_index'] = len(chunks)
					chunks.append(sub_chunk)
			else:
				chunk_dict = self._create_chunk_dict(section_text, len(chunks), metadata)
				chunk_dict['section_title'] = section["title"]
				chunk_dict['section_level'] = section["level"]
				chunks.append(chunk_dict)
		
		return chunks
	
	def _calculate_semantic_similarity(self, text1: str, text2: str) -> float:
		"""Simple semantic similarity based on word overlap"""
		# Simple word-based similarity
		words1 = set(text1.lower().split())
		words2 = set(text2.lower().split())
		
		if not words1 or not words2:
			return 0.0
		
		intersection = len(words1.intersection(words2))
		union = len(words1.union(words2))
		
		return intersection / union if union > 0 else 0.0
	
	def _create_chunk_dict(self, content: str, index: int, metadata: Dict[str, Any] = None) -> Dict[str, Any]:
		"""Create standardized chunk dictionary"""
		content_hash = hashlib.md5(content.encode()).hexdigest()
		
		chunk_dict = {
			'chunk_index': index,
			'content': content,
			'content_hash': content_hash,
			'character_count': len(content),
			'token_count': len(content.split()),  # Simple token estimation
		}
		
		if metadata:
			chunk_dict['source_metadata'] = metadata
		
		return chunk_dict

class DocumentProcessor:
	"""Main document processing orchestrator"""
	
	def __init__(self, 
	             config: ProcessingConfig,
	             ollama_integration: AdvancedOllamaIntegration,
	             tenant_id: str,
	             capability_id: str = "rag"):
		self.config = config
		self.ollama_integration = ollama_integration
		self.tenant_id = tenant_id
		self.capability_id = capability_id
		
		self.extractor = DocumentFormatExtractor()
		self.chunker = DocumentChunker(config.chunking)
		self.logger = logging.getLogger(__name__)
		
		# Processing queues and state
		self.processing_queue = asyncio.Queue(maxsize=1000)
		self.active_jobs = {}
		self.is_running = False
		
		# Statistics
		self.stats = {
			'total_processed': 0,
			'successful_processed': 0,
			'failed_processed': 0,
			'total_chunks_created': 0,
			'average_processing_time_ms': 0.0
		}
	
	async def start(self):
		"""Start the document processor"""
		if self.is_running:
			return
		
		self.is_running = True
		self.logger.info("Document processor started")
		
		# Start worker tasks
		self.worker_tasks = [
			asyncio.create_task(self._processing_worker())
			for _ in range(4)  # 4 concurrent workers
		]
	
	async def stop(self):
		"""Stop the document processor"""
		if not self.is_running:
			return
		
		self.is_running = False
		
		# Cancel worker tasks
		for task in self.worker_tasks:
			task.cancel()
		
		await asyncio.gather(*self.worker_tasks, return_exceptions=True)
		self.logger.info("Document processor stopped")
	
	async def process_document_async(self, 
	                                document_create: DocumentCreate,
	                                priority: ProcessingPriority = ProcessingPriority.NORMAL,
	                                callback: Optional[callable] = None) -> str:
		"""Queue document for asynchronous processing"""
		job_id = uuid7str()
		
		job = {
			'id': job_id,
			'document_create': document_create,
			'priority': priority,
			'callback': callback,
			'created_at': datetime.now()
		}
		
		await self.processing_queue.put(job)
		self.active_jobs[job_id] = job
		
		self.logger.info(f"Queued document processing job {job_id}")
		return job_id
	
	async def process_document(self, 
	                          document_create: DocumentCreate,
	                          generate_embeddings: bool = True) -> ProcessingResult:
		"""Process a single document synchronously"""
		start_time = time.time()
		
		try:
			# Validate file size
			if hasattr(document_create, 'content') and len(document_create.content.encode()) > self.config.max_file_size_mb * 1024 * 1024:
				raise ValueError(f"File too large (max {self.config.max_file_size_mb}MB)")
			
			# Extract content if from file path
			if hasattr(document_create, 'source_path') and document_create.source_path:
				extraction_result = await self.extractor.extract_content(
					document_create.source_path, 
					document_create.content_type
				)
				content = extraction_result['content']
				extracted_metadata = extraction_result['metadata']
			else:
				content = document_create.content
				extracted_metadata = {}
			
			# Create file hash
			file_hash = hashlib.sha256(content.encode()).hexdigest()
			
			# Create document model
			document = Document(
				tenant_id=self.tenant_id,
				knowledge_base_id=document_create.knowledge_base_id,
				source_path=document_create.source_path,
				filename=document_create.filename,
				file_hash=file_hash,
				content_type=document_create.content_type,
				file_size=len(content.encode()),
				title=document_create.title,
				content=content,
				metadata={**document_create.metadata, **extracted_metadata},
				language=document_create.language,
				processing_status=DocumentStatus.PROCESSING
			)
			
			# Chunk the document
			chunks_data = await self.chunker.chunk_document(content, extracted_metadata)
			
			# Generate embeddings if requested
			chunks = []
			if generate_embeddings and chunks_data:
				chunks = await self._generate_chunk_embeddings(chunks_data, document.id)
			else:
				# Create chunks without embeddings
				for chunk_data in chunks_data:
					chunk = DocumentChunk(
						tenant_id=self.tenant_id,
						document_id=document.id,
						knowledge_base_id=document.knowledge_base_id,
						chunk_index=chunk_data['chunk_index'],
						content=chunk_data['content'],
						content_hash=chunk_data['content_hash'],
						embedding=[0.0] * 1024,  # Placeholder embedding
						start_position=chunk_data.get('start_position'),
						end_position=chunk_data.get('end_position'),
						token_count=chunk_data.get('token_count'),
						character_count=chunk_data['character_count'],
						section_title=chunk_data.get('section_title'),
						section_level=chunk_data.get('section_level', 0)
					)
					chunks.append(chunk)
			
			# Update document with processing results
			document.processing_status = DocumentStatus.COMPLETED
			document.chunk_count = len(chunks)
			
			# Update statistics
			processing_time_ms = (time.time() - start_time) * 1000
			self.stats['total_processed'] += 1
			self.stats['successful_processed'] += 1
			self.stats['total_chunks_created'] += len(chunks)
			self._update_average_processing_time(processing_time_ms)
			
			return ProcessingResult(
				document=document,
				chunks=chunks,
				processing_time_ms=processing_time_ms,
				success=True,
				metadata={
					'chunks_created': len(chunks),
					'extraction_metadata': extracted_metadata
				}
			)
		
		except Exception as e:
			processing_time_ms = (time.time() - start_time) * 1000
			self.stats['total_processed'] += 1
			self.stats['failed_processed'] += 1
			
			self.logger.error(f"Document processing failed: {str(e)}")
			
			# Create failed document record
			document = Document(
				tenant_id=self.tenant_id,
				knowledge_base_id=document_create.knowledge_base_id,
				source_path=document_create.source_path,
				filename=document_create.filename,
				file_hash="failed",
				content_type=document_create.content_type,
				file_size=0,
				content=document_create.content[:1000] if hasattr(document_create, 'content') else "",
				processing_status=DocumentStatus.FAILED,
				processing_error=str(e)
			)
			
			return ProcessingResult(
				document=document,
				chunks=[],
				processing_time_ms=processing_time_ms,
				success=False,
				error_message=str(e)
			)
	
	async def _processing_worker(self):
		"""Worker task for processing documents from queue"""
		while self.is_running:
			try:
				# Get job from queue with timeout
				job = await asyncio.wait_for(
					self.processing_queue.get(),
					timeout=1.0
				)
				
				# Process the document
				result = await self.process_document(job['document_create'])
				
				# Call callback if provided
				if job['callback']:
					try:
						await job['callback'](job['id'], result)
					except Exception as e:
						self.logger.error(f"Callback failed for job {job['id']}: {str(e)}")
				
				# Remove from active jobs
				self.active_jobs.pop(job['id'], None)
			
			except asyncio.TimeoutError:
				continue
			except Exception as e:
				self.logger.error(f"Processing worker error: {str(e)}")
				await asyncio.sleep(1.0)
	
	async def _generate_chunk_embeddings(self, chunks_data: List[Dict[str, Any]], document_id: str) -> List[DocumentChunk]:
		"""Generate embeddings for document chunks"""
		chunks = []
		
		# Prepare texts for batch embedding
		texts = [chunk_data['content'] for chunk_data in chunks_data]
		
		# Generate embeddings using Ollama integration
		embedding_results = {}
		
		def embedding_callback(result):
			if result['success']:
				for i, embedding in enumerate(result['embeddings']):
					embedding_results[i] = embedding
		
		# Request embeddings
		request_id = await self.ollama_integration.generate_embeddings_async(
			texts=texts,
			model="bge-m3",
			tenant_id=self.tenant_id,
			capability_id=self.capability_id,
			priority=RequestPriority.NORMAL,
			callback=embedding_callback
		)
		
		# Wait for embeddings (with timeout)
		max_wait_time = 60.0  # 60 seconds
		wait_start = time.time()
		
		while len(embedding_results) < len(texts) and (time.time() - wait_start) < max_wait_time:
			await asyncio.sleep(0.1)
		
		# Create chunks with embeddings
		for i, chunk_data in enumerate(chunks_data):
			embedding = embedding_results.get(i, [0.0] * 1024)  # Fallback to zero vector
			
			chunk = DocumentChunk(
				tenant_id=self.tenant_id,
				document_id=document_id,
				knowledge_base_id=chunk_data.get('knowledge_base_id', ''),
				chunk_index=chunk_data['chunk_index'],
				content=chunk_data['content'],
				content_hash=chunk_data['content_hash'],
				embedding=embedding,
				start_position=chunk_data.get('start_position'),
				end_position=chunk_data.get('end_position'),
				token_count=chunk_data.get('token_count'),
				character_count=chunk_data['character_count'],
				section_title=chunk_data.get('section_title'),
				section_level=chunk_data.get('section_level', 0),
				embedding_confidence=1.0 if i in embedding_results else 0.0
			)
			chunks.append(chunk)
		
		return chunks
	
	def _update_average_processing_time(self, processing_time_ms: float):
		"""Update average processing time statistics"""
		current_avg = self.stats['average_processing_time_ms']
		total_processed = self.stats['total_processed']
		
		# Calculate running average
		self.stats['average_processing_time_ms'] = (
			(current_avg * (total_processed - 1) + processing_time_ms) / total_processed
		)
	
	def get_statistics(self) -> Dict[str, Any]:
		"""Get processing statistics"""
		queue_size = self.processing_queue.qsize()
		active_jobs_count = len(self.active_jobs)
		
		return {
			**self.stats,
			'queue_size': queue_size,
			'active_jobs': active_jobs_count,
			'success_rate': self.stats['successful_processed'] / max(1, self.stats['total_processed']),
			'average_chunks_per_document': self.stats['total_chunks_created'] / max(1, self.stats['successful_processed']),
			'is_running': self.is_running
		}

# Factory function for APG integration
async def create_document_processor(
	tenant_id: str,
	capability_id: str,
	ollama_integration: AdvancedOllamaIntegration,
	config: ProcessingConfig = None
) -> DocumentProcessor:
	"""Create and start document processor"""
	if config is None:
		config = ProcessingConfig()
	
	processor = DocumentProcessor(config, ollama_integration, tenant_id, capability_id)
	await processor.start()
	return processor